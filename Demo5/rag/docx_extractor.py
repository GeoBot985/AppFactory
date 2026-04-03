from docx import Document
import os

def extract_docx_text(path: str) -> str:
    """
    Extracts text from a .docx file including paragraphs and simple tables.
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")

    try:
        doc = Document(path)
    except Exception as e:
        raise ValueError(f"Failed to read DOCX file.")

    full_text = []

    # Get all body elements in the document
    # doc.paragraphs only gives paragraphs (not those inside tables)
    # doc.tables only gives tables
    # However, for a simple implementation as requested:
    # "extract paragraph text in order"
    # "Minimum acceptable extraction: body paragraphs from the main document"
    # "If easy with current library, include simple table text extraction"

    # Let's iterate through the document paragraphs and then tables as separate blocks.

    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            full_text.append(stripped)

    # Simple table extraction
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            full_text.append(" | ".join(row_text))

    return "\n\n".join(full_text)
